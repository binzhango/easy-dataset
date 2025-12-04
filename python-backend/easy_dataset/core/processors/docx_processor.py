"""
DOCX file processor implementation.

This module provides DOCX (Microsoft Word) document processing using python-docx.
It extracts text with formatting, preserves paragraph structure, and handles
tables and lists.
"""

from pathlib import Path
from typing import Dict, List

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError(
        "python-docx is required for DOCX processing. "
        "Install it with: pip install python-docx"
    )

from ..file_processor import FileProcessor, FileType, ProcessedDocument


class DOCXProcessor(FileProcessor):
    """
    Processor for DOCX (Microsoft Word) documents.
    
    Extracts text with formatting preservation, maintains paragraph structure,
    and handles tables and lists. Uses python-docx for document parsing.
    """
    
    def supports(self, file_type: FileType) -> bool:
        """Check if this processor supports DOCX files."""
        return file_type == FileType.DOCX
    
    def process(self, file_path: Path) -> ProcessedDocument:
        """
        Process a DOCX file and extract its content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            ProcessedDocument containing text, structure, and metadata
            
        Raises:
            ValueError: If DOCX is invalid or corrupted
            IOError: If file cannot be read
        """
        self.validate_file(file_path)
        
        try:
            # Open the document
            doc = Document(file_path)
            
            # Extract text content with structure
            text_content = self._extract_text(doc)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            return ProcessedDocument(
                content=text_content,
                metadata=metadata,
                images=[]
            )
            
        except Exception as e:
            raise IOError(f"Error reading DOCX file: {e}")
    
    def _extract_text(self, doc: Document) -> str:
        """
        Extract text from document with structure preservation.
        
        This method iterates through all document elements (paragraphs and tables)
        in order, preserving the document structure.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            Extracted text with structure preserved
        """
        text_parts = []
        
        # Iterate through document body elements
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # It's a paragraph
                paragraph = Paragraph(element, doc)
                para_text = self._extract_paragraph_text(paragraph)
                if para_text:
                    text_parts.append(para_text)
                    
            elif isinstance(element, CT_Tbl):
                # It's a table
                table = Table(element, doc)
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)
        
        return '\n\n'.join(text_parts)
    
    def _extract_paragraph_text(self, paragraph: Paragraph) -> str:
        """
        Extract text from a paragraph.
        
        Args:
            paragraph: python-docx Paragraph instance
            
        Returns:
            Paragraph text
        """
        text = paragraph.text.strip()
        
        # Add formatting indicators for headings
        if paragraph.style.name.startswith('Heading'):
            # Extract heading level
            try:
                level = int(paragraph.style.name.split()[-1])
                # Add markdown-style heading markers
                text = f"{'#' * level} {text}"
            except (ValueError, IndexError):
                pass
        
        return text
    
    def _extract_table_text(self, table: Table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: python-docx Table instance
            
        Returns:
            Table text in markdown-like format
        """
        table_lines = []
        
        for i, row in enumerate(table.rows):
            # Extract cell text
            cells = [cell.text.strip() for cell in row.cells]
            
            # Format as markdown table
            row_text = ' | '.join(cells)
            table_lines.append(f"| {row_text} |")
            
            # Add separator after header row
            if i == 0:
                separator = ' | '.join(['---'] * len(cells))
                table_lines.append(f"| {separator} |")
        
        return '\n'.join(table_lines)
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """
        Extract metadata from DOCX document.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            Dictionary containing document metadata
        """
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
        }
        
        # Extract core properties if available
        if doc.core_properties:
            props = doc.core_properties
            
            # Add available properties
            if props.title:
                metadata['title'] = props.title
            if props.author:
                metadata['author'] = props.author
            if props.subject:
                metadata['subject'] = props.subject
            if props.keywords:
                metadata['keywords'] = props.keywords
            if props.created:
                metadata['created'] = props.created.isoformat()
            if props.modified:
                metadata['modified'] = props.modified.isoformat()
            if props.last_modified_by:
                metadata['last_modified_by'] = props.last_modified_by
        
        # Count headings
        heading_count = sum(
            1 for para in doc.paragraphs
            if para.style.name.startswith('Heading')
        )
        metadata['heading_count'] = heading_count
        
        return metadata
    
    def extract_paragraphs(self, doc: Document) -> List[Dict]:
        """
        Extract paragraphs with their styles.
        
        Useful for more detailed document analysis.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            List of paragraph dictionaries with text and style
        """
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    'style': para.style.name,
                    'is_heading': para.style.name.startswith('Heading'),
                })
        
        return paragraphs
    
    def extract_tables(self, doc: Document) -> List[List[List[str]]]:
        """
        Extract tables as structured data.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            List of tables, each table is a list of rows,
            each row is a list of cell texts
        """
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return tables
